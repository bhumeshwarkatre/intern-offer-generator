import os
import re
import csv
import qrcode
import random
import string
import tempfile
import streamlit as st
from datetime import date
from docxtpl import DocxTemplate
from docx.shared import Inches
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from smtplib import SMTP
import mammoth
import imgkit
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

# --- Config ---
st.set_page_config("Intern Offer Generator", layout="wide")
EMAIL = st.secrets["email"]["user"]
PASSWORD = st.secrets["email"]["password"]
TEMPLATE_FILE = "temp_offer_letter.docx"
CSV_FILE = "intern_offers.csv"
LOGO = "logo.png"

# --- Styles ---
st.markdown("""
<style>
.title-text { font-size: 2rem; font-weight: 700; }
.stButton>button { background-color: #1E88E5; color: white; padding: 0.5rem 1.5rem; border-radius: 8px; font-weight: 600; }
button[title="View fullscreen"] { visibility: hidden !important; }
</style>
""", unsafe_allow_html=True)

# --- Header ---
with st.container():
    col1, col2 = st.columns([1, 6])
    with col1:
        if os.path.exists(LOGO):
            st.image(LOGO, width=80)
    with col2:
        st.markdown('<div class="title-text">SkyHighes Technologies Internship Letter Portal</div>', unsafe_allow_html=True)

st.divider()

# --- Utility Functions ---
def format_date(d): return d.strftime("%A, %d %B %Y")

def generate_certificate_key(): return ''.join(random.choices(string.ascii_uppercase + string.digits, k=9))

def generate_qr(data):
    qr = qrcode.QRCode(box_size=10, border=4)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    path = os.path.join(tempfile.gettempdir(), "qr.png")
    img.save(path)
    return path

def save_to_csv(data):
    exists = os.path.exists(CSV_FILE)
    with open(CSV_FILE, mode='a', newline='') as f:
        writer = csv.writer(f)
        if not exists: writer.writerow(data.keys())
        writer.writerow(data.values())

def convert_docx_to_png_pdf(docx_path):
    # Convert .docx → HTML
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value

    html_file = docx_path.replace(".docx", ".html")
    with open(html_file, "w", encoding="utf-8") as f:
        f.write(html)

    # HTML → PNG using imgkit
    png_file = docx_path.replace(".docx", ".png")
    imgkit.from_file(html_file, png_file)

    # PNG → PDF using reportlab
    final_pdf = docx_path.replace(".docx", ".pdf")
    c = canvas.Canvas(final_pdf, pagesize=A4)
    c.drawImage(png_file, 0, 0, width=A4[0], height=A4[1])
    c.save()

    return final_pdf

def send_email(receiver, pdf_path, data):
    msg = MIMEMultipart()
    msg["From"] = EMAIL
    msg["To"] = receiver
    msg["Subject"] = f"🎉 Your Internship Offer - {data['intern_name']}"

    html = f"""<p>Dear {data['intern_name']},</p>
    <p>We’re pleased to offer you an internship at <b>SkyHighes Technology</b>.</p>
    <p>Your offer letter is attached.</p>"""

    msg.attach(MIMEText(html, "html"))

    with open(pdf_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(pdf_path)}")
        msg.attach(part)

    with SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(EMAIL, PASSWORD)
        server.send_message(msg)

# --- Form ---
with st.form("offer_form"):
    col1, col2, col3 = st.columns(3)
    with col1: intern_name = st.text_input("Intern Name", placeholder="e.g., Bhumeshwar Katre")
    with col2: domain = st.text_input("Domain", placeholder="e.g., Web Development")
    with col3: email = st.text_input("Recipient Email", placeholder="e.g., intern@example.com")

    col4, col5, col6 = st.columns(3)
    with col4: start_date = st.date_input("Start Date", value=date.today())
    with col5: end_date = st.date_input("End Date", value=date.today())
    with col6: offer_date = st.date_input("Offer Date", value=date.today())

    submit = st.form_submit_button("🚀 Generate & Send Offer Letter")

# --- On Submit ---
if submit:
    if not all([intern_name, domain, email]):
        st.error("❌ Please fill all fields.")
    elif not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        st.warning("⚠️ Invalid email.")
    elif end_date < start_date:
        st.warning("⚠️ End date cannot be before start date.")
    elif not os.path.exists(TEMPLATE_FILE):
        st.error("❌ Offer template missing.")
    else:
        intern_id = generate_certificate_key()
        data = {
            "intern_name": intern_name.strip(),
            "domain": domain.strip(),
            "start_date": format_date(start_date),
            "end_date": format_date(end_date),
            "offer_date": format_date(offer_date),
            "i_id": intern_id,
            "email": email.strip()
        }

        save_to_csv(data)
        doc = DocxTemplate(TEMPLATE_FILE)
        doc.render(data)

        qr_path = generate_qr(f"{intern_name}, {domain}, {start_date}, {end_date}, {offer_date}, {intern_id}")
        try:
            doc.tables[0].rows[0].cells[2].paragraphs[0].add_run().add_picture(qr_path, width=Inches(1.5))
        except:
            st.warning("⚠️ QR insertion failed.")

        docx_path = os.path.join(tempfile.gettempdir(), f"Offer_{intern_name}.docx")
        doc.save(docx_path)

        try:
            pdf_path = convert_docx_to_png_pdf(docx_path)
            send_email(email, pdf_path, data)
            st.success(f"✅ Sent to {email}")
            with open(pdf_path, "rb") as f:
                st.download_button("📥 Download Offer Letter", f, file_name=os.path.basename(pdf_path))
        except Exception as e:
            st.error(f"❌ Failed to send: {e}")
